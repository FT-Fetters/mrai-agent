import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
# from docx.enum.section import WD_ORIENTATION, WD_SECTION # Not used
from docx.oxml.ns import qn
# from docx.oxml import OxmlElement # Not used directly now
# import docx.opc.constants # Not used
from docx.shared import Twips
from loguru import logger
import os
from typing import List, Dict, Optional

from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
# Removed redundant: from docx.oxml.shared import qn (Ensuring this is removed)

from mrai.agent.schema import Tool

# Configure logging (No longer needed with loguru defaults)
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

def word_tool_list():
    return [
        ReadWordTool(),
        CreateWordTool(),
        AddParagraphTool(),
        AddTableTool(),
        ModifyParagraphTool(),
        ModifyTableCellTool(),
        ApplyRunFormattingTool(),
        ApplyParagraphFormattingTool(),
        InsertParagraphTool(),
        DeleteParagraphTool(),
        DeleteTableTool(),
    ]

class ReadWordTool(Tool):
    """A tool to read content and basic formatting from Word documents."""

    def __init__(self):
        super().__init__(
            name="read_word",
            description="Reads content and basic formatting from a Word document, including paragraphs, text formatting, and tables.",
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path",
                    type="string",
                    description="The path to the Word document file.",
                    required=True
                )
            }
        )

    def _format_value(self, value, unit="", precision=1):
        """Safely formats a numeric value, handling None."""
        if value is None:
            return "Default"
        # Ensure value is numeric before formatting
        try:
            # Attempt conversion to float for formatting
            float_value = float(value)
            return f"{float_value:.{precision}f}{unit}"
        except (ValueError, TypeError):
            # If conversion fails, return the original value as string
            return str(value)

    def _get_run_formatting(self, run):
        """Extracts formatting information for a Run."""
        font = run.font
        formatting = {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_name": font.name,
            # Handle potential AttributeError if size is None
            "font_size_pt": font.size.pt if font.size and hasattr(font.size, 'pt') else None,
            # Handle potential AttributeError if color or rgb is None
            "font_color_rgb": font.color.rgb if font.color and hasattr(font.color, 'rgb') else None,
        }
        return formatting

    def _get_paragraph_formatting(self, paragraph):
        """Extracts formatting information for a Paragraph."""
        para_format = paragraph.paragraph_format
        # Helper to safely access attributes that might be None
        def safe_get_attr(obj, attr, unit):
            val = getattr(obj, attr, None)
            return val.cm if val and hasattr(val, 'cm') and unit == 'cm' else (val.pt if val and hasattr(val, 'pt') and unit == 'pt' else val)

        formatting = {
            "alignment": para_format.alignment,
            "line_spacing": para_format.line_spacing,
            "line_spacing_rule": para_format.line_spacing_rule,
            "first_line_indent_cm": safe_get_attr(para_format, 'first_line_indent', 'cm'),
            "left_indent_cm": safe_get_attr(para_format, 'left_indent', 'cm'),
            "right_indent_cm": safe_get_attr(para_format, 'right_indent', 'cm'),
            "space_before_pt": safe_get_attr(para_format, 'space_before', 'pt'),
            "space_after_pt": safe_get_attr(para_format, 'space_after', 'pt'),
        }
        return formatting

    def _process_paragraph(self, para, index):
        """Processes a single paragraph and returns its formatted string."""
        if not para.text.strip(): # Skip empty paragraphs
            return None

        output = []
        output.append(f"\n### Paragraph {index}")
        output.append(f"**Text:** {para.text}")

        para_format = self._get_paragraph_formatting(para)
        align_str = str(para_format['alignment']) if para_format['alignment'] else 'Default'
        spacing_str = self._format_value(para_format['line_spacing'], "", 1)
        spacing_rule_str = str(para_format['line_spacing_rule']) if para_format['line_spacing_rule'] else 'Default'
        first_indent_str = self._format_value(para_format['first_line_indent_cm'], "cm", 2)
        left_indent_str = self._format_value(para_format['left_indent_cm'], "cm", 2)
        right_indent_str = self._format_value(para_format['right_indent_cm'], "cm", 2)
        space_before_str = self._format_value(para_format['space_before_pt'], "pt", 1)
        space_after_str = self._format_value(para_format['space_after_pt'], "pt", 1)

        output.append(f"**Paragraph Formatting:** Alignment={align_str}, Line Spacing={spacing_str} ({spacing_rule_str}), "
                      f"First Line Indent={first_indent_str}, Left Indent={left_indent_str}, Right Indent={right_indent_str}, "
                      f"Space Before={space_before_str}, Space After={space_after_str}")

        if para.runs:
            output.append("**Run Formatting:**")
            for run in para.runs:
                if not run.text.strip(): # Skip empty runs
                    continue
                run_format = self._get_run_formatting(run)

                rgb_color = run_format['font_color_rgb']
                color_str = 'Default'
                if isinstance(rgb_color, RGBColor): # Check if it's an RGBColor object
                    try:
                        # Format as two-digit hex
                        color_str = f"#{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}"
                    except (TypeError, IndexError, ValueError):
                        logger.warning(f"Could not format RGB color value: {rgb_color}. Using 'Default'.")
                        # Keep color_str as 'Default'
                elif rgb_color is not None: # Handle cases where it might be something else unexpected
                     logger.warning(f"Unexpected color format: {rgb_color}. Using 'Default'.")


                size_str = self._format_value(run_format['font_size_pt'], "pt", 1)
                font_name_str = run_format['font_name'] if run_format['font_name'] else 'Default'
                output.append(f"- `{run_format['text']}`: "
                              f"Font='{font_name_str}', "
                              f"Size={size_str}, "
                              f"Color={color_str}, "
                              f"Bold={'Yes' if run_format['bold'] else 'No'}, "
                              f"Italic={'Yes' if run_format['italic'] else 'No'}, "
                              f"Underline={'Yes' if run_format['underline'] else 'No'}")
        output.append("\n---")
        return "\n".join(output)

    def _process_table(self, table, index):
        """Processes a single table and returns its formatted Markdown string."""
        output = []
        output.append(f"\n### Table {index}")

        # Robustly get column count
        num_cols = 0
        if hasattr(table, 'columns') and table.columns:
             num_cols = len(table.columns)
        elif table.rows and table.rows[0].cells: # Fallback: Use first row's cell count if exists
            num_cols = len(table.rows[0].cells)

        # Try to get table style name
        style_name = "Default"
        if hasattr(table, 'style') and table.style and hasattr(table.style, 'name'):
            style_name = table.style.name
        output.append(f"Table Style: '{style_name}'")

        try:
            tblPr = table._tbl.tblPr
            if tblPr is not None:
                # Safely find the tblBorders element
                tblBorders_element = tblPr.find(qn('w:tblBorders'))
                if tblBorders_element is not None:
                    # Check for common border types by looking for child elements
                    has_borders = any(
                        tblBorders_element.find(qn(f'w:{tag}')) is not None
                        for tag in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']
                    )
                    if has_borders:
                        output.append("Table Borders: Defined (specific styles vary)")
        except Exception as e:
            logger.debug(f"Could not check table-level borders for table {index}: {e}")
        # >>> END ADDED

        output.append(f"(Approximate Columns: {num_cols}, Rows: {len(table.rows)})\n")

        formatting_notes = [] # Store formatting notes (row_idx, col_idx_or_0, note)

        if num_cols > 0 and table.rows:
            self._add_row_formatting_notes(table.rows[0], 1, formatting_notes)

            # Build Markdown table header from first row
            header_cells_text = []
            first_row_cells = table.rows[0].cells
            effective_num_cols_header = 0 # Track columns considering spans in header
            col_idx_markdown_header = 0 # Track markdown column index for header
            for i in range(len(first_row_cells)):
                 # Check if we already filled this markdown column due to a previous span
                 if col_idx_markdown_header >= len(header_cells_text):
                     cell = first_row_cells[i]
                     # Extract text
                     cell_text = " ".join(p.text for p in cell.paragraphs).strip().replace("|", "\\|")
                     header_cells_text.append(cell_text)

                     # Check for horizontal span
                     grid_span = 1
                     tcPr = cell._tc.tcPr # Get cell properties element
                     if tcPr is not None:
                         grid_span_elem = tcPr.gridSpan
                         if grid_span_elem is not None and grid_span_elem.val is not None:
                             try:
                                 grid_span = int(grid_span_elem.val)
                             except (ValueError, TypeError):
                                 grid_span = 1

                     # Add formatting notes (use 1-based index for notes)
                     self._add_cell_formatting_notes(cell, 1, col_idx_markdown_header + 1, formatting_notes)
                     if grid_span > 1:
                         formatting_notes.append((1, col_idx_markdown_header + 1, f"Horizontally spans {grid_span} columns"))
                         # Add placeholder cells for Markdown structure
                         header_cells_text.extend([""] * (grid_span - 1))
                         col_idx_markdown_header += grid_span
                     else:
                         col_idx_markdown_header += 1
                 else:
                     # This cell position is covered by a previous span, skip the actual cell
                     # but advance markdown header index if needed (though span handling should cover this)
                     # This logic might need refinement if complex spans exist
                     pass

            # Adjust num_cols based on the effective width found in the header
            num_cols = max(num_cols, len(header_cells_text))

            # Pad header if needed (less likely now but safety)
            while len(header_cells_text) < num_cols:
                 header_cells_text.append(f"Col_{len(header_cells_text)+1}")

            header = "| " + " | ".join(header_cells_text) + " |"
            separator = "|-" + "-|".join(['-' * max(3, len(str(txt))) for txt in header_cells_text]) + "-|"

            output.append(header)
            output.append(separator)

            # Fill table content (iterate through rows *starting from the second row*)
            for row_idx, row in enumerate(table.rows[1:], start=1): # Start from row index 1 (second row)
                self._add_row_formatting_notes(row, row_idx + 1, formatting_notes)

                row_content = []
                cells_to_process = row.cells
                col_idx_actual = 0 # Index for accessing cells_to_process
                col_idx_markdown = 0 # Index for markdown columns, accounting for spans

                while col_idx_markdown < num_cols and col_idx_actual < len(cells_to_process):
                    # Ensure the target markdown column isn't already filled by a previous span
                    if col_idx_markdown >= len(row_content):
                        cell = cells_to_process[col_idx_actual]
                        cell_text = "\n".join([p.text for p in cell.paragraphs])
                        escaped_text = cell_text.replace("|", "\\|").strip()
                        escaped_text = escaped_text.replace('\n', '<br>')
                        row_content.append(escaped_text)

                        # Check for horizontal span
                        grid_span = 1
                        tcPr = cell._tc.tcPr
                        if tcPr is not None:
                             grid_span_elem = tcPr.gridSpan
                             if grid_span_elem is not None and grid_span_elem.val is not None:
                                 try:
                                     grid_span = int(grid_span_elem.val)
                                 except (ValueError, TypeError):
                                     grid_span = 1

                        # Add formatting notes (use row_idx+1 because enumerate starts at 1 here, col_idx_markdown+1)
                        self._add_cell_formatting_notes(cell, row_idx + 1, col_idx_markdown + 1, formatting_notes)
                        if grid_span > 1:
                            formatting_notes.append((row_idx + 1, col_idx_markdown + 1, f"Horizontally spans {grid_span} columns"))
                            # Add placeholder cells for Markdown structure
                            row_content.extend([""] * (grid_span - 1))
                            col_idx_markdown += grid_span
                        else:
                            col_idx_markdown += 1
                    else:
                        # This markdown column was filled by a span from a cell earlier in this row.
                        # We still need to process the *next* actual cell.
                        pass # Let the outer loop increment col_idx_markdown implicitly

                    col_idx_actual += 1

                # Pad row if it has fewer cells than num_cols (due to spans or short rows)
                while len(row_content) < num_cols:
                    row_content.append("")
                    # col_idx_markdown += 1 # Already handled by loop structure

                output.append("| " + " | ".join(row_content) + " |")

        elif table.rows: # Has rows but couldn't determine columns reliably
             output.append("*(Table found, but structure unclear - printing raw row content)*")
             for r_idx, row in enumerate(table.rows):
                 row_text = " | ".join([" ".join(p.text for p in cell.paragraphs).strip() for cell in row.cells])
                 output.append(f"Row {r_idx+1}: {row_text}")
        else:
            output.append("*(Empty Table)*")


        # Add formatting notes if any were collected
        if formatting_notes:
             output.append("\n**Formatting Notes:**")
             # Sort notes primarily by row, then by column (0 for row notes first)
             formatting_notes.sort(key=lambda x: (x[0], x[1]))
             for r, c, note in formatting_notes:
                 if c == 0: # Row-specific note
                     output.append(f"*   Row {r}: {note}")
                 else: # Cell-specific note
                      output.append(f"*   Row {r}, Col {c}: {note}")


        output.append("\n---")
        return "\n".join(output)

    def _add_row_formatting_notes(self, row, row_idx_1_based, notes):
        """Helper to extract and add row height formatting notes."""
        try:
            height = row.height
            height_rule = row.height_rule

            rule_str = None
            if height_rule is not None and height_rule != WD_ROW_HEIGHT_RULE.AUTO:
                # Get the rule name from the enum value
                try:
                    rule_str = WD_ROW_HEIGHT_RULE(height_rule).name
                except ValueError:
                    rule_str = str(height_rule) # Fallback to raw value

            height_str = None
            if height is not None:
                 # height is in Twips, convert to points
                 height_pt = height.pt # Direct conversion using .pt attribute
                 height_str = f"{height_pt:.1f}pt"


            # Add note only if there's non-default info
            if height_str or rule_str:
                note_parts = []
                if height_str:
                    note_parts.append(f"Height={height_str}")
                if rule_str:
                    note_parts.append(f"Rule={rule_str}")
                # Use column 0 to signify a row-level note
                notes.append((row_idx_1_based, 0, ", ".join(note_parts)))

        except Exception as e:
            logger.debug(f"Could not get row height/rule for row index {row_idx_1_based}: {e}")

    def _add_cell_formatting_notes(self, cell, row_idx, col_idx, notes):
        """Helper to extract and add cell formatting notes."""
        tcPr = cell._tc.tcPr # Get cell properties element once

        # --- Text Formatting (First Run) ---
        try:
            if cell.paragraphs and cell.paragraphs[0].runs:
                first_run = cell.paragraphs[0].runs[0]
                if first_run.text.strip(): # Only if the run has text
                    run_format = self._get_run_formatting(first_run)
                    # Format details concisely
                    size_str = self._format_value(run_format['font_size_pt'], "pt", 1)
                    font_name_str = run_format['font_name'] if run_format['font_name'] else 'Default'

                    rgb_color = run_format['font_color_rgb']
                    color_str = 'Default'
                    if isinstance(rgb_color, RGBColor):
                        try:
                            color_str = f"#{rgb_color[0]:02X}{rgb_color[1]:02X}{rgb_color[2]:02X}"
                        except (TypeError, IndexError, ValueError): pass # Keep default on error
                    elif rgb_color is not None: pass # Keep default for unexpected

                    format_parts = [f"Font='{font_name_str}'", f"Size={size_str}"]
                    if color_str != 'Default': format_parts.append(f"Color={color_str}")
                    if run_format['bold']: format_parts.append("Bold")
                    if run_format['italic']: format_parts.append("Italic")
                    if run_format['underline']: format_parts.append("Underline")
                    notes.append((row_idx, col_idx, f"Text Format (Start): {', '.join(format_parts)}"))
        except Exception as e:
            logger.debug(f"Could not get cell text formatting for R{row_idx}C{col_idx}: {e}")


        # --- Horizontal Alignment ---
        try:
            if cell.paragraphs:
                first_para_align = cell.paragraphs[0].alignment
                if first_para_align is not None:
                    align_str = str(first_para_align).split('.')[-1]
                    if align_str != 'LEFT': # LEFT is usually default
                        notes.append((row_idx, col_idx, f"Alignment={align_str}"))
        except Exception as e:
             logger.debug(f"Could not get cell paragraph alignment for R{row_idx}C{col_idx}: {e}")


        # --- Vertical Alignment ---
        try:
            # Access vertical alignment via cell properties (tcPr -> vAlign)
            # tcPr = cell._tc.tcPr # Already defined above
            v_align = None
            if tcPr is not None:
                v_align_elem = tcPr.vAlign
                if v_align_elem is not None and v_align_elem.val is not None:
                    v_align = v_align_elem.val # This is usually the enum value like 'center', 'top', etc.

            is_top = False
            if isinstance(v_align, int): # Check if it's the integer enum value
                is_top = (v_align == WD_ALIGN_VERTICAL.TOP)
            elif isinstance(v_align, str): # Check if it's the string value
                is_top = (v_align.lower() == 'top')

            if v_align is not None and not is_top:
                 valign_str = str(v_align).upper() # Get string representation
                 notes.append((row_idx, col_idx, f"Vertical Alignment={valign_str}"))
        except Exception as e:
             logger.debug(f"Could not get cell vertical alignment for R{row_idx}C{col_idx}: {e}")

        # --- Cell Width ---
        try:
             if tcPr is not None:
                 tcW = tcPr.tcW # Get width element
                 if tcW is not None and tcW.w is not None:
                     width_val = tcW.w
                     width_type = tcW.type
                     if width_type == 'dxa': # Twentieths of a point
                         width_pt = width_val / 20.0
                         notes.append((row_idx, col_idx, f"Width={width_pt:.1f}pt"))
                     elif width_type == 'pct': # Percentage
                         notes.append((row_idx, col_idx, f"Width={width_val}%"))
                     elif width_type == 'auto' or width_type is None:
                         notes.append((row_idx, col_idx, "Width=Auto"))
                     else: # Other types?
                          notes.append((row_idx, col_idx, f"Width={width_val} ({width_type})"))
        except Exception as e:
              logger.debug(f"Could not get cell width for R{row_idx}C{col_idx}: {e}")

        # --- Cell Borders ---
        try:
             if tcPr is not None:
                 # Safely find the tcBorders element
                 tcBorders_element = tcPr.find(qn('w:tcBorders'))
                 if tcBorders_element is not None:
                     # Check if any specific border element exists
                     has_specific_borders = any(
                         tcBorders_element.find(qn(f'w:{tag}')) is not None
                         # We just check for existence, value check might be too complex here
                         for tag in ['top', 'bottom', 'left', 'right'] # Others like tl2br? maybe keep simple
                             # Note: insideH/V might not be typical on tcBorders, mostly on tblBorders
                     )
                     if has_specific_borders:
                         notes.append((row_idx, col_idx, "Has Specific Borders"))
        except Exception as e:
              logger.debug(f"Could not check cell borders for R{row_idx}C{col_idx}: {e}")

         # Note: Detecting shading requires deeper XML parsing (tcPr -> shd)

    def execute(self, file_path: str) -> str:
        """
        Reads the content and formatting information from a Word document,
        processing paragraphs and tables in their original order.

        Args:
            file_path (str): The path to the Word document file.

        Returns:
            str: A string containing the document content and formatting information,
                 formatted using Markdown.
        """
        try:
            document = Document(file_path)
            output = [f"# Word Document Content: {file_path}\n"]
            para_idx = 0
            table_idx = 0

            # Iterate through the document body elements to maintain order
            # Parent element access requires understanding the structure; docx simplifies this.
            # We rely on the assumption that document.paragraphs and document.tables
            # hold references to the objects in the order they appear relative to their type.
            # To get absolute order, we iterate through the XML body's children.
            parent_element = document.element.body
            if parent_element is None:
                 logger.error("Could not access document body element.")
                 return "Error: Could not access document body."

            for child in parent_element:
                # Check if the element is a paragraph ('w:p') or a table ('w:tbl')
                if child.tag == qn('w:p'):
                    if para_idx < len(document.paragraphs):
                        para = document.paragraphs[para_idx]
                        para_output = self._process_paragraph(para, para_idx + 1)
                        if para_output: # Only add if not empty/skipped
                             output.append(para_output)
                        para_idx += 1
                    else:
                         # This might happen if there are paragraph elements not captured by document.paragraphs (e.g., in headers/footers if not handled)
                         logger.warning(f"Found a paragraph element (index {para_idx}) beyond the count in document.paragraphs ({len(document.paragraphs)}). Skipping.")


                elif child.tag == qn('w:tbl'):
                    if table_idx < len(document.tables):
                        table = document.tables[table_idx]
                        table_output = self._process_table(table, table_idx + 1)
                        output.append(table_output)
                        table_idx += 1
                    else:
                         logger.warning(f"Found a table element (index {table_idx}) beyond the count in document.tables ({len(document.tables)}). Skipping.")
                # Note: This loop might not capture elements in headers, footers, text boxes etc.
                # It focuses on the main document body's flow.


            # Simple validation check - might log warnings if counts don't match
            if para_idx != len(document.paragraphs):
                logger.warning(f"Processed {para_idx} paragraphs based on body iteration, but document.paragraphs contains {len(document.paragraphs)}.")
            if table_idx != len(document.tables):
                 logger.warning(f"Processed {table_idx} tables based on body iteration, but document.tables contains {len(document.tables)}.")


            return "\n".join(output)

        except FileNotFoundError:
            logger.error(f"Error: File not found - {file_path}")
            return f"Error: File not found - {file_path}"
        except ImportError as e:
             logger.error(f"Import error, likely missing dependency: {e}. Please ensure 'python-docx' and 'loguru' are installed.")
             return f"Import error: {e}. Make sure 'python-docx' and 'loguru' are installed."
        except Exception as e:
            # Log the full traceback for better debugging using loguru
            logger.exception(f"An unexpected error occurred while reading Word file '{file_path}': {e}")
            return f"An unexpected error occurred while reading Word file '{file_path}': {e}"



class CreateWordTool(Tool):
    """Creates a new, empty Word document at the specified path."""
    def __init__(self):
        super().__init__(
            name="create_word_document",
            description="Creates a new, empty Word document (.docx). If the file already exists, it will be overwritten.",
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path",
                    type="string",
                    description="The full path where the new Word document should be saved (e.g., /path/to/new_document.docx).",
                    required=True
                )
            }
        )

    def execute(self, file_path: str) -> str:
        """
        Creates a new Word document.

        Args:
            file_path (str): The path to save the new document.

        Returns:
            str: A confirmation message or an error message.
        """
        try:
            # Ensure the directory exists
            dir_path = os.path.dirname(file_path)
            if dir_path and not os.path.exists(dir_path):
                os.makedirs(dir_path)
                logger.info(f"Created directory: {dir_path}")

            document = Document()
            document.save(file_path)
            logger.info(f"Successfully created new Word document at: {file_path}")
            return f"Successfully created new Word document at: {file_path}"
        except Exception as e:
            logger.exception(f"Error creating Word document at '{file_path}': {e}")
            return f"Error creating Word document at '{file_path}': {e}"

class AddParagraphTool(Tool):
    """Adds a paragraph to the end of a Word document."""
    def __init__(self):
        super().__init__(
            name="add_paragraph_to_word",
            description="Adds a new paragraph with the specified text to the end of an existing Word document. Optionally applies a built-in style.",
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path",
                    type="string",
                    description="The path to the Word document file to modify.",
                    required=True
                ),
                "text": Tool.ToolParameter(
                    name="text",
                    type="string",
                    description="The text content of the new paragraph.",
                    required=True
                ),
                "style": Tool.ToolParameter(
                    name="style",
                    type="string",
                    description="Optional. The name of a built-in Word style to apply (e.g., 'Heading 1', 'Body Text', 'Normal'). If omitted or invalid, default paragraph style is used.",
                    required=False
                )
            }
        )

    def execute(self, file_path: str, text: str, style: Optional[str] = None) -> str:
        """
        Adds a paragraph to a Word document.

        Args:
            file_path (str): Path to the document.
            text (str): Text to add.
            style (Optional[str]): Style name to apply.

        Returns:
            str: Confirmation or error message.
        """
        try:
            if not os.path.exists(file_path):
                return f"Error: File not found at '{file_path}'. Use create_word_document first if needed."

            document = Document(file_path)

            # Add the paragraph with optional style
            paragraph = document.add_paragraph(text)
            if style:
                try:
                    # Check if the style exists in the document's known styles
                    # Note: This is a basic check; truly validating requires more complex checks or handling potential errors.
                    # python-docx doesn't provide a direct way to list all available built-in style names easily without introspection.
                    paragraph.style = document.styles[style] # type: ignore
                    style_applied_msg = f" with style '{style}'"
                except KeyError:
                    logger.warning(f"Style '{style}' not found in document styles. Using default paragraph style.")
                    style_applied_msg = f" (Warning: Style '{style}' not found, used default)"
                except Exception as style_e: # Catch other potential style errors
                    logger.warning(f"Error applying style '{style}': {style_e}. Using default paragraph style.")
                    style_applied_msg = f" (Warning: Error applying style '{style}', used default)"
            else:
                style_applied_msg = ""


            document.save(file_path)
            logger.info(f"Successfully added paragraph to '{file_path}'{style_applied_msg}.")
            return f"Successfully added paragraph to '{file_path}'{style_applied_msg}."

        except Exception as e:
            logger.exception(f"Error adding paragraph to Word document '{file_path}': {e}")
            return f"Error adding paragraph to Word document '{file_path}': {e}"


class AddTableTool(Tool):
    """Adds a table to the end of a Word document."""
    def __init__(self):
        super().__init__(
            name="add_table_to_word",
            description="Adds a new table with specified dimensions to the end of an existing Word document. Optionally applies a built-in table style and populates the first row as a header.",
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path",
                    type="string",
                    description="The path to the Word document file to modify.",
                    required=True
                ),
                "rows": Tool.ToolParameter(
                    name="rows",
                    type="number",
                    description="The number of rows for the new table.",
                    required=True
                ),
                "cols": Tool.ToolParameter(
                    name="cols",
                    type="number",
                    description="The number of columns for the new table.",
                    required=True
                ),
                "style": Tool.ToolParameter(
                    name="style",
                    type="string",
                    description="Optional. The name of a built-in Word table style to apply (e.g., 'Table Grid', 'Light Shading Accent 1'). If omitted or invalid, the default table style is used.",
                    required=False
                ),
                 "header_row": Tool.ToolParameter(
                    name="header_row",
                    type="list",
                    description="Optional. A list of strings representing the content for the first row (header). The list length must match the number of columns.",
                    required=False
                )
            }
        )

    def execute(self, file_path: str, rows: int, cols: int, style: Optional[str] = None, header_row: Optional[List[str]] = None) -> str:
        """
        Adds a table to a Word document.

        Args:
            file_path (str): Path to the document.
            rows (int): Number of rows.
            cols (int): Number of columns.
            style (Optional[str]): Table style name.
            header_row (Optional[List[str]]): List of header cell contents.

        Returns:
            str: Confirmation or error message.
        """
        try:
            if not os.path.exists(file_path):
                 return f"Error: File not found at '{file_path}'. Use create_word_document first if needed."
            if rows <= 0 or cols <= 0:
                 return "Error: Number of rows and columns must be positive."
            if header_row and len(header_row) != cols:
                return f"Error: Header row list length ({len(header_row)}) does not match the number of columns ({cols})."


            document = Document(file_path)

            # Add the table
            table = document.add_table(rows=rows, cols=cols)
            style_applied_msg = ""

            # Apply style if specified
            if style:
                try:
                    table.style = document.styles[style] # type: ignore
                    style_applied_msg = f" with style '{style}'"
                except KeyError:
                    logger.warning(f"Table style '{style}' not found in document styles. Using default table style.")
                    style_applied_msg = f" (Warning: Table Style '{style}' not found, used default)"
                except Exception as style_e:
                    logger.warning(f"Error applying table style '{style}': {style_e}. Using default table style.")
                    style_applied_msg = f" (Warning: Error applying table style '{style}', used default)"

            # Populate header row if provided
            header_populated_msg = ""
            if header_row:
                hdr_cells = table.rows[0].cells
                for i, header_text in enumerate(header_row):
                    hdr_cells[i].text = header_text
                header_populated_msg = " and populated header row"


            # Add a blank paragraph after the table for spacing, unless it's the very last element
            # This often improves layout when viewing the document.
            document.add_paragraph()


            document.save(file_path)
            logger.info(f"Successfully added {rows}x{cols} table to '{file_path}'{style_applied_msg}{header_populated_msg}.")
            return f"Successfully added {rows}x{cols} table to '{file_path}'{style_applied_msg}{header_populated_msg}."

        except Exception as e:
            logger.exception(f"Error adding table to Word document '{file_path}': {e}")
            return f"Error adding table to Word document '{file_path}': {e}"



class ModifyParagraphTool(Tool):
    """Modifies a specific paragraph in a Word document."""
    def __init__(self):
        super().__init__(
            name="modify_paragraph_in_word",
            description="Modifies the text or style of a specific paragraph in a Word document. The paragraph can be identified either by its 1-based index or by unique text it contains. At least one modification (new_text or new_style) must be provided.",
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path",
                    type="string",
                    description="The path to the Word document file to modify.",
                    required=True
                ),
                "paragraph_identifier": Tool.ToolParameter(
                    name="paragraph_identifier",
                    type="string",
                    description="Identifier for the paragraph. Can be its 1-based index (e.g., '5') or a unique string of text contained within the paragraph (e.g., 'report introduction').",
                    required=True
                ),
                "new_text": Tool.ToolParameter(
                    name="new_text",
                    type="string",
                    description="Optional. The new text content to replace the paragraph's current text. If omitted, the text is not changed.",
                    required=False
                ),
                "new_style": Tool.ToolParameter(
                    name="new_style",
                    type="string",
                    description="Optional. The name of a built-in Word style to apply to the paragraph (e.g., 'Heading 1', 'Body Text'). If omitted, the style is not changed.",
                    required=False
                )
            }
        )

    def execute(self, file_path: str, paragraph_identifier: str, new_text: Optional[str] = None, new_style: Optional[str] = None) -> str:
        """
        Modifies a paragraph's text or style.

        Args:
            file_path (str): Path to the document.
            paragraph_identifier (str): Index (as string) or text content to find the paragraph.
            new_text (Optional[str]): New text for the paragraph.
            new_style (Optional[str]): New style name for the paragraph.

        Returns:
            str: Confirmation or error message.
        """
        if new_text is None and new_style is None:
            return "Error: You must provide either 'new_text' or 'new_style' to modify the paragraph."

        try:
            if not os.path.exists(file_path):
                return f"Error: File not found at '{file_path}'."

            document = Document(file_path)
            target_paragraph = None
            identifier_type = "unknown"

            # Try identifying by index first
            try:
                para_index_1_based = int(paragraph_identifier)
                if 1 <= para_index_1_based <= len(document.paragraphs):
                    target_paragraph = document.paragraphs[para_index_1_based - 1]
                    identifier_type = f"index {para_index_1_based}"
                else:
                    return f"Error: Paragraph index {para_index_1_based} is out of range (1 to {len(document.paragraphs)})."
            except ValueError:
                # If not an integer, assume it's text content
                identifier_type = f"text '{paragraph_identifier}'"
                found = False
                for i, para in enumerate(document.paragraphs):
                    if paragraph_identifier in para.text:
                        target_paragraph = para
                        logger.info(f"Found paragraph containing text '{paragraph_identifier}' at index {i+1}.")
                        found = True
                        break
                if not found:
                    return f"Error: Could not find any paragraph containing the text: '{paragraph_identifier}'."

            if target_paragraph is None: # Should technically be caught above, but safety check
                 return f"Error: Could not identify the paragraph using '{paragraph_identifier}'."

            modified_parts = []

            # Apply new style if provided
            if new_style:
                try:
                    target_paragraph.style = document.styles[new_style] # type: ignore
                    modified_parts.append(f"style set to '{new_style}'")
                    logger.info(f"Applied style '{new_style}' to paragraph {identifier_type}.")
                except KeyError:
                    logger.warning(f"Style '{new_style}' not found. Style was not changed.")
                    # Optionally return an error here, or just log and continue
                    return f"Error: Style '{new_style}' not found in the document. No changes made to style."
                except Exception as style_e:
                    logger.warning(f"Error applying style '{new_style}': {style_e}. Style was not changed.")
                    return f"Error applying style '{new_style}': {style_e}. No changes made to style."

            # Replace text if provided
            if new_text is not None:
                # --- Try to preserve formatting from the first run --- START
                original_font = None
                if target_paragraph.runs:
                    original_font = target_paragraph.runs[0].font
                # --- Try to preserve formatting from the first run --- END

                # Clear existing content (runs) within the paragraph
                # Based on https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
                p_element = target_paragraph._element
                p_element.clear_content()
                # Add the new text as a single run
                new_run = target_paragraph.add_run(new_text)

                # --- Apply preserved formatting --- START
                if original_font:
                    new_run.font.name = original_font.name
                    new_run.font.size = original_font.size
                    new_run.font.bold = original_font.bold
                    new_run.font.italic = original_font.italic
                    new_run.font.underline = original_font.underline
                    new_run.font.color.rgb = original_font.color.rgb
                    # Copy other attributes as needed (e.g., strike, small_caps etc.)
                # --- Apply preserved formatting --- END

                modified_parts.append("text updated (attempted style preservation)")
                logger.info(f"Replaced text in paragraph {identifier_type}, attempted style preservation.")

            document.save(file_path)
            modification_summary = " and ".join(modified_parts)
            return f"Successfully modified paragraph identified by {identifier_type}: {modification_summary} in '{file_path}'."

        except Exception as e:
            logger.exception(f"Error modifying paragraph in Word document '{file_path}': {e}")
            return f"Error modifying paragraph in Word document '{file_path}': {e}"


class ModifyTableCellTool(Tool):
    """Modifies the text content of a specific cell in a table within a Word document."""
    def __init__(self):
        super().__init__(
            name="modify_table_cell_in_word",
            description="Modifies the text content of a specific cell within a table in a Word document. The table and cell are identified by their 1-based indices.",
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path",
                    type="string",
                    description="The path to the Word document file to modify.",
                    required=True
                ),
                "table_index": Tool.ToolParameter(
                    name="table_index",
                    type="number",
                    description="The 1-based index of the table within the document.",
                    required=True
                ),
                "row_index": Tool.ToolParameter(
                    name="row_index",
                    type="number",
                    description="The 1-based index of the row within the table.",
                    required=True
                ),
                "col_index": Tool.ToolParameter(
                    name="col_index",
                    type="number",
                    description="The 1-based index of the column within the row.",
                    required=True
                ),
                "new_text": Tool.ToolParameter(
                    name="new_text",
                    type="string",
                    description="The new text content for the specified cell.",
                    required=True
                )
            }
        )

    def execute(self, file_path: str, table_index: int, row_index: int, col_index: int, new_text: str) -> str:
        """
        Modifies the text in a table cell.

        Args:
            file_path (str): Path to the document.
            table_index (int): 1-based index of the table.
            row_index (int): 1-based index of the row.
            col_index (int): 1-based index of the column.
            new_text (str): New text for the cell.

        Returns:
            str: Confirmation or error message.
        """
        try:
            if not os.path.exists(file_path):
                return f"Error: File not found at '{file_path}'."

            document = Document(file_path)

            # Validate table index
            if not (1 <= table_index <= len(document.tables)):
                return f"Error: Table index {table_index} is out of range (1 to {len(document.tables)})."
            table = document.tables[table_index - 1]

            # Validate row index
            if not (1 <= row_index <= len(table.rows)):
                return f"Error: Row index {row_index} is out of range (1 to {len(table.rows)}) for table {table_index}."
            row = table.rows[row_index - 1]

            # Validate column index
            if not (1 <= col_index <= len(row.cells)):
                 # Note: This uses the actual cell count in the specific row, which can vary due to merged cells.
                 # A more robust approach might consider the table's intended column count if merged cells are common.
                 return f"Error: Column index {col_index} is out of range (1 to {len(row.cells)}) for table {table_index}, row {row_index}."
            cell = row.cells[col_index - 1]

            # --- Try to preserve formatting from the first run of the first paragraph --- START
            original_font = None
            first_paragraph = None
            if cell.paragraphs:
                first_paragraph = cell.paragraphs[0]
                if first_paragraph.runs:
                    original_font = first_paragraph.runs[0].font
            # --- Try to preserve formatting --- END

            # Clear cell content by replacing text in the first paragraph
            # More robust clearing might iterate through paragraphs and clear runs if needed
            cell.text = "" # Clear text, this might remove paragraphs other than the first

            # Ensure there is at least one paragraph to add the run to
            if not cell.paragraphs:
                 cell.add_paragraph("") # Add an empty paragraph if cleared completely
            target_para_in_cell = cell.paragraphs[0]

            # Add new text as a run
            new_run = target_para_in_cell.add_run(new_text)

            # --- Apply preserved formatting --- START
            if original_font:
                new_run.font.name = original_font.name
                new_run.font.size = original_font.size
                new_run.font.bold = original_font.bold
                new_run.font.italic = original_font.italic
                new_run.font.underline = original_font.underline
                new_run.font.color.rgb = original_font.color.rgb
                # Copy other relevant font attributes if necessary
            # --- Apply preserved formatting --- END

            document.save(file_path)
            logger.info(f"Successfully modified cell ({row_index}, {col_index}) in table {table_index} in '{file_path}', attempted style preservation.")
            return f"Successfully modified cell at Table {table_index}, Row {row_index}, Col {col_index} to '{new_text}' in '{file_path}', attempted style preservation."

        except Exception as e:
            logger.exception(f"Error modifying table cell in Word document '{file_path}': {e}")
            return f"Error modifying table cell in Word document '{file_path}': {e}"

class ApplyRunFormattingTool(Tool):
    """Applies formatting to specific text runs within a paragraph."""
    def __init__(self):
        super().__init__(
            name="apply_run_formatting_in_word",
            description=(
                "Finds specific text within a target paragraph (identified by index or contained text) "
                "and applies formatting (bold, italic, underline, font size, font name, color) to the text run(s) containing it. "
                "Specify at least one formatting option."
            ),
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path", type="string",
                    description="Path to the Word document.", required=True
                ),
                "paragraph_identifier": Tool.ToolParameter(
                    name="paragraph_identifier", type="string",
                    description="Identifier for the target paragraph (1-based index or unique contained text).", required=True
                ),
                "target_text": Tool.ToolParameter(
                    name="target_text", type="string",
                    description="The specific text within the paragraph to apply formatting to. The formatting will be applied to the entire run containing this text.", required=True
                ),
                "bold": Tool.ToolParameter(
                    name="bold", type="boolean",
                    description="Apply bold formatting.", required=False
                ),
                "italic": Tool.ToolParameter(
                    name="italic", type="boolean",
                    description="Apply italic formatting.", required=False
                ),
                "underline": Tool.ToolParameter(
                    name="underline", type="boolean",
                    description="Apply underline formatting.", required=False
                ),
                "font_size_pt": Tool.ToolParameter(
                    name="font_size_pt", type="number",
                    description="Set font size in points (e.g., 12).", required=False
                ),
                "font_name": Tool.ToolParameter(
                    name="font_name", type="string",
                    description="Set font name (e.g., 'Calibri', 'Times New Roman').", required=False
                ),
                "font_color_rgb": Tool.ToolParameter(
                    name="font_color_rgb", type="string",
                    description="Set font color as a 6-digit hex RGB string (e.g., 'FF0000' for red).", required=False
                )
            }
        )

    def execute(self, file_path: str, paragraph_identifier: str, target_text: str,
                  bold: Optional[bool] = None, italic: Optional[bool] = None, underline: Optional[bool] = None,
                  font_size_pt: Optional[float] = None, font_name: Optional[str] = None, font_color_rgb: Optional[str] = None) -> str:

        formatting_options = [bold, italic, underline, font_size_pt, font_name, font_color_rgb]
        if all(opt is None for opt in formatting_options):
            return "Error: At least one formatting option (bold, italic, underline, font_size_pt, font_name, font_color_rgb) must be provided."

        try:
            if not os.path.exists(file_path):
                return f"Error: File not found at '{file_path}'."

            document = Document(file_path)
            target_paragraph = None
            identifier_type = "unknown"
            found_paragraph = False

            # Find the paragraph (copied logic from ModifyParagraphTool)
            try:
                para_index_1_based = int(paragraph_identifier)
                if 1 <= para_index_1_based <= len(document.paragraphs):
                    target_paragraph = document.paragraphs[para_index_1_based - 1]
                    identifier_type = f"index {para_index_1_based}"
                    found_paragraph = True
                else:
                    return f"Error: Paragraph index {para_index_1_based} is out of range (1 to {len(document.paragraphs)})."
            except ValueError:
                identifier_type = f"text '{paragraph_identifier}'"
                for i, para in enumerate(document.paragraphs):
                    if paragraph_identifier in para.text:
                        target_paragraph = para
                        identifier_type = f"text '{paragraph_identifier}' (found at index {i+1})"
                        found_paragraph = True
                        logger.info(f"Found paragraph containing text '{paragraph_identifier}' at index {i+1}.")
                        break
                if not found_paragraph:
                    return f"Error: Could not find paragraph containing text: '{paragraph_identifier}'."

            if not found_paragraph or target_paragraph is None:
                return f"Error: Could not identify the paragraph using '{paragraph_identifier}'."

            # Find runs containing the target text and apply formatting
            runs_modified_count = 0
            applied_formats = []
            for run in target_paragraph.runs:
                if target_text in run.text:
                    font = run.font
                    if bold is not None: font.bold = bold; applied_formats.append(f"bold={bold}")
                    if italic is not None: font.italic = italic; applied_formats.append(f"italic={italic}")
                    if underline is not None: font.underline = underline; applied_formats.append(f"underline={underline}")
                    if font_size_pt is not None: font.size = Pt(font_size_pt); applied_formats.append(f"size={font_size_pt}pt")
                    if font_name is not None: font.name = font_name; applied_formats.append(f"font='{font_name}'")
                    if font_color_rgb is not None:
                        try:
                            # Basic validation for 6-digit hex
                            if len(font_color_rgb) == 6 and all(c in '0123456789abcdefABCDEF' for c in font_color_rgb):
                                font.color.rgb = RGBColor.from_string(font_color_rgb)
                                applied_formats.append(f"color=#{font_color_rgb}")
                            else:
                                logger.warning(f"Invalid RGB color format: '{font_color_rgb}'. Must be 6-digit hex. Color not applied.")
                                # Decide whether to error out or just skip color
                                # return f"Error: Invalid RGB color format: '{font_color_rgb}'. Must be 6-digit hex."
                        except Exception as color_e:
                             logger.warning(f"Error applying RGB color '{font_color_rgb}': {color_e}. Color not applied.")
                    runs_modified_count += 1
                    logger.info(f"Applied formatting to run containing '{target_text}' in paragraph {identifier_type}.")
                    # Note: This applies to the *entire run*. Splitting runs for partial text is complex.

            if runs_modified_count == 0:
                return f"Error: Text '{target_text}' not found within the identified paragraph ({identifier_type}). No formatting applied."

            document.save(file_path)
            format_summary = ", ".join(sorted(list(set(applied_formats)))) # Unique formats applied
            return f"Successfully applied formatting ({format_summary}) to {runs_modified_count} run(s) containing '{target_text}' in paragraph {identifier_type} in '{file_path}'."

        except Exception as e:
            logger.exception(f"Error applying run formatting in Word document '{file_path}': {e}")
            return f"Error applying run formatting in Word document '{file_path}': {e}"

class ApplyParagraphFormattingTool(Tool):
    """Applies uniform formatting to all runs within a specific paragraph."""
    def __init__(self):
        super().__init__(
            name="apply_paragraph_formatting_in_word",
            description=(
                "Applies uniform formatting (bold, italic, underline, font size, font name, color) "
                "to all text runs within a target paragraph (identified by index or contained text). "
                "This affects the entire paragraph's text. Specify at least one formatting option."
            ),
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path", type="string",
                    description="Path to the Word document.", required=True
                ),
                "paragraph_identifier": Tool.ToolParameter(
                    name="paragraph_identifier", type="string",
                    description="Identifier for the target paragraph (1-based index or unique contained text).", required=True
                ),
                # Formatting options are the same as ApplyRunFormattingTool
                "bold": Tool.ToolParameter(
                    name="bold", type="boolean",
                    description="Apply bold formatting to the entire paragraph.", required=False
                ),
                "italic": Tool.ToolParameter(
                    name="italic", type="boolean",
                    description="Apply italic formatting to the entire paragraph.", required=False
                ),
                "underline": Tool.ToolParameter(
                    name="underline", type="boolean",
                    description="Apply underline formatting to the entire paragraph.", required=False
                ),
                "font_size_pt": Tool.ToolParameter(
                    name="font_size_pt", type="number",
                    description="Set font size in points for the entire paragraph.", required=False
                ),
                "font_name": Tool.ToolParameter(
                    name="font_name", type="string",
                    description="Set font name for the entire paragraph.", required=False
                ),
                "font_color_rgb": Tool.ToolParameter(
                    name="font_color_rgb", type="string",
                    description="Set font color (6-digit hex RGB) for the entire paragraph.", required=False
                )
            }
        )

    def execute(self, file_path: str, paragraph_identifier: str,
                  bold: Optional[bool] = None, italic: Optional[bool] = None, underline: Optional[bool] = None,
                  font_size_pt: Optional[float] = None, font_name: Optional[str] = None, font_color_rgb: Optional[str] = None) -> str:

        formatting_options = [bold, italic, underline, font_size_pt, font_name, font_color_rgb]
        if all(opt is None for opt in formatting_options):
            return "Error: At least one formatting option must be provided."

        try:
            if not os.path.exists(file_path):
                return f"Error: File not found at '{file_path}'."

            document = Document(file_path)
            target_paragraph = None
            identifier_type = "unknown"
            found_paragraph = False

            # Find the paragraph (copied logic from ModifyParagraphTool/ApplyRunFormattingTool)
            try:
                para_index_1_based = int(paragraph_identifier)
                if 1 <= para_index_1_based <= len(document.paragraphs):
                    target_paragraph = document.paragraphs[para_index_1_based - 1]
                    identifier_type = f"index {para_index_1_based}"
                    found_paragraph = True
                else:
                    return f"Error: Paragraph index {para_index_1_based} is out of range (1 to {len(document.paragraphs)})."
            except ValueError:
                identifier_type = f"text '{paragraph_identifier}'"
                for i, para in enumerate(document.paragraphs):
                    if paragraph_identifier in para.text:
                        target_paragraph = para
                        identifier_type = f"text '{paragraph_identifier}' (found at index {i+1})"
                        found_paragraph = True
                        logger.info(f"Found paragraph containing text '{paragraph_identifier}' at index {i+1}.")
                        break
                if not found_paragraph:
                    return f"Error: Could not find paragraph containing text: '{paragraph_identifier}'."

            if not found_paragraph or target_paragraph is None:
                return f"Error: Could not identify the paragraph using '{paragraph_identifier}'."

            # Apply formatting to all runs in the paragraph
            applied_formats = []
            if not target_paragraph.runs:
                 logger.warning(f"Paragraph {identifier_type} has no runs (is empty?). Cannot apply formatting.")
                 # Or potentially add an empty run and format it? Decided against it for now.

            for run in target_paragraph.runs:
                font = run.font
                # Apply requested formatting unconditionally to each run
                if bold is not None: font.bold = bold; applied_formats.append(f"bold={bold}")
                if italic is not None: font.italic = italic; applied_formats.append(f"italic={italic}")
                if underline is not None: font.underline = underline; applied_formats.append(f"underline={underline}")
                if font_size_pt is not None: font.size = Pt(font_size_pt); applied_formats.append(f"size={font_size_pt}pt")
                if font_name is not None: font.name = font_name; applied_formats.append(f"font='{font_name}'")
                if font_color_rgb is not None:
                    try:
                        if len(font_color_rgb) == 6 and all(c in '0123456789abcdefABCDEF' for c in font_color_rgb):
                            font.color.rgb = RGBColor.from_string(font_color_rgb)
                            applied_formats.append(f"color=#{font_color_rgb}")
                        else:
                            logger.warning(f"Invalid RGB color format: '{font_color_rgb}'. Color not applied.")
                            # Consider returning error if strictness is needed
                    except Exception as color_e:
                         logger.warning(f"Error applying RGB color '{font_color_rgb}': {color_e}. Color not applied to this run.")

            document.save(file_path)
            format_summary = ", ".join(sorted(list(set(applied_formats)))) # Unique formats applied
            if not applied_formats: # Handle case where formatting failed (e.g., bad color format)
                return f"Attempted to apply formatting to paragraph {identifier_type}, but no valid formats were specified or applied successfully in '{file_path}'."

            logger.info(f"Applied uniform formatting ({format_summary}) to all runs in paragraph {identifier_type}.")
            return f"Successfully applied uniform formatting ({format_summary}) to paragraph {identifier_type} in '{file_path}'."

        except Exception as e:
            logger.exception(f"Error applying paragraph formatting in Word document '{file_path}': {e}")
            return f"Error applying paragraph formatting in Word document '{file_path}': {e}"

class InsertParagraphTool(Tool):
    """Inserts a new paragraph before or after a specified paragraph."""
    def __init__(self):
        super().__init__(
            name="insert_paragraph_in_word",
            description="Inserts a new paragraph with specified text either before or after a target paragraph (identified by index or contained text).",
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path", type="string",
                    description="Path to the Word document.", required=True
                ),
                "target_paragraph_identifier": Tool.ToolParameter(
                    name="target_paragraph_identifier", type="string",
                    description="Identifier for the target paragraph (1-based index or unique contained text) relative to which the new paragraph will be inserted.", required=True
                ),
                "text_to_insert": Tool.ToolParameter(
                    name="text_to_insert", type="string",
                    description="The text content of the new paragraph to insert.", required=True
                ),
                "insert_before": Tool.ToolParameter(
                    name="insert_before", type="boolean",
                    description="Set to true to insert before the target paragraph, false (or omit) to insert after.", required=False
                ),
                 "style": Tool.ToolParameter(
                    name="style", type="string",
                    description="Optional. Style for the new paragraph (e.g., 'Heading 1').", required=False
                )
            }
        )

    def execute(self, file_path: str, target_paragraph_identifier: str, text_to_insert: str,
                  insert_before: bool = False, style: Optional[str] = None) -> str:
        try:
            if not os.path.exists(file_path):
                return f"Error: File not found at '{file_path}'."

            document = Document(file_path)
            target_paragraph = None
            target_para_element = None
            identifier_type = "unknown"
            found_paragraph = False

            # Find the paragraph and its element (slightly modified find logic)
            try:
                para_index_1_based = int(target_paragraph_identifier)
                if 1 <= para_index_1_based <= len(document.paragraphs):
                    target_paragraph = document.paragraphs[para_index_1_based - 1]
                    target_para_element = target_paragraph._element
                    identifier_type = f"index {para_index_1_based}"
                    found_paragraph = True
                else:
                    return f"Error: Target paragraph index {para_index_1_based} is out of range (1 to {len(document.paragraphs)})."
            except ValueError:
                identifier_type = f"text '{target_paragraph_identifier}'"
                for i, para in enumerate(document.paragraphs):
                    if target_paragraph_identifier in para.text:
                        target_paragraph = para
                        target_para_element = para._element
                        identifier_type = f"text '{target_paragraph_identifier}' (found at index {i+1})"
                        found_paragraph = True
                        logger.info(f"Found target paragraph containing text '{target_paragraph_identifier}' at index {i+1}.")
                        break
                if not found_paragraph:
                    return f"Error: Could not find target paragraph containing text: '{target_paragraph_identifier}'."

            if not found_paragraph or target_paragraph is None or target_para_element is None:
                return f"Error: Could not identify the target paragraph using '{target_paragraph_identifier}'."

            # Create the new paragraph element
            new_para = document.add_paragraph(text_to_insert) # Add temporarily to end
            style_applied_msg = ""
            if style:
                try:
                    new_para.style = document.styles[style] # type: ignore
                    style_applied_msg = f" with style '{style}'"
                except KeyError:
                    logger.warning(f"Style '{style}' not found, using default.")
                    style_applied_msg = " (style not found, used default)"
                except Exception as style_e:
                    logger.warning(f"Error applying style '{style}': {style_e}. Using default.")
                    style_applied_msg = f" (error applying style '{style}', used default)"

            new_para_element = new_para._element # Get the OxmlElement of the new paragraph
            # Remove the new paragraph from the end where add_paragraph placed it
            new_para_element.getparent().remove(new_para_element)

            # Insert the new paragraph element at the correct position
            if insert_before:
                target_para_element.addprevious(new_para_element)
                position_desc = "before"
            else:
                target_para_element.addnext(new_para_element)
                position_desc = "after"

            document.save(file_path)
            logger.info(f"Inserted paragraph '{text_to_insert}' {position_desc} paragraph {identifier_type}.")
            return f"Successfully inserted paragraph '{text_to_insert}'{style_applied_msg} {position_desc} paragraph identified by {identifier_type} in '{file_path}'."

        except Exception as e:
            logger.exception(f"Error inserting paragraph in Word document '{file_path}': {e}")
            return f"Error inserting paragraph in Word document '{file_path}': {e}"

class DeleteParagraphTool(Tool):
    """Deletes a specific paragraph from a Word document."""
    def __init__(self):
        super().__init__(
            name="delete_paragraph_in_word",
            description="Deletes a specific paragraph identified either by its 1-based index or by unique text it contains.",
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path", type="string",
                    description="Path to the Word document.", required=True
                ),
                "paragraph_identifier": Tool.ToolParameter(
                    name="paragraph_identifier", type="string",
                    description="Identifier for the paragraph to delete (1-based index or unique contained text).", required=True
                )
            }
        )

    def execute(self, file_path: str, paragraph_identifier: str) -> str:
        try:
            if not os.path.exists(file_path):
                return f"Error: File not found at '{file_path}'."

            document = Document(file_path)
            target_paragraph = None
            target_para_element = None
            identifier_type = "unknown"
            found_paragraph = False
            target_index = -1

            # Find the paragraph and its element
            try:
                para_index_1_based = int(paragraph_identifier)
                if 1 <= para_index_1_based <= len(document.paragraphs):
                    target_index = para_index_1_based -1
                    target_paragraph = document.paragraphs[target_index]
                    target_para_element = target_paragraph._element
                    identifier_type = f"index {para_index_1_based}"
                    found_paragraph = True
                else:
                    return f"Error: Paragraph index {para_index_1_based} is out of range (1 to {len(document.paragraphs)})."
            except ValueError:
                identifier_type = f"text '{paragraph_identifier}'"
                for i, para in enumerate(document.paragraphs):
                    if paragraph_identifier in para.text:
                        target_index = i
                        target_paragraph = para
                        target_para_element = para._element
                        identifier_type = f"text '{paragraph_identifier}' (found at index {i+1})"
                        found_paragraph = True
                        logger.info(f"Found paragraph containing text '{paragraph_identifier}' at index {i+1}.")
                        break
                if not found_paragraph:
                    return f"Error: Could not find paragraph containing text: '{paragraph_identifier}'."

            if not found_paragraph or target_paragraph is None or target_para_element is None:
                 return f"Error: Could not identify the paragraph to delete using '{paragraph_identifier}'."

            # Delete the paragraph element from its parent
            parent_element = target_para_element.getparent()
            if parent_element is not None:
                parent_element.remove(target_para_element)
                # Note: This removes the paragraph from the XML. The document.paragraphs list
                # might not immediately reflect this change until the document is reloaded.
                # However, saving the document persists the deletion.
                document.save(file_path)
                logger.info(f"Deleted paragraph identified by {identifier_type} (original index {target_index+1}).")
                return f"Successfully deleted paragraph identified by {identifier_type} from '{file_path}'."
            else:
                # This case should be rare for paragraphs in the main body
                logger.error(f"Could not find parent element for paragraph {identifier_type} to delete.")
                return f"Error: Could not delete paragraph {identifier_type} - parent element not found."

        except Exception as e:
            logger.exception(f"Error deleting paragraph in Word document '{file_path}': {e}")
            return f"Error deleting paragraph in Word document '{file_path}': {e}"

class DeleteTableTool(Tool):
    """Deletes a specific table from a Word document."""
    def __init__(self):
        super().__init__(
            name="delete_table_in_word",
            description="Deletes a specific table identified by its 1-based index.",
            parameters={
                "file_path": Tool.ToolParameter(
                    name="file_path", type="string",
                    description="Path to the Word document.", required=True
                ),
                "table_index": Tool.ToolParameter(
                    name="table_index", type="number",
                    description="The 1-based index of the table to delete.", required=True
                )
            }
        )

    def execute(self, file_path: str, table_index: int) -> str:
        try:
            if not os.path.exists(file_path):
                return f"Error: File not found at '{file_path}'."

            document = Document(file_path)

            # Validate table index
            if not (1 <= table_index <= len(document.tables)):
                return f"Error: Table index {table_index} is out of range (1 to {len(document.tables)})."

            table_to_delete = document.tables[table_index - 1]
            table_element = table_to_delete._element

            # Delete the table element from its parent
            parent_element = table_element.getparent()
            if parent_element is not None:
                parent_element.remove(table_element)
                # Similar to paragraph deletion, saving persists the change.
                document.save(file_path)
                logger.info(f"Deleted table at index {table_index} from '{file_path}'.")
                return f"Successfully deleted table at index {table_index} from '{file_path}'."
            else:
                 logger.error(f"Could not find parent element for table index {table_index} to delete.")
                 return f"Error: Could not delete table at index {table_index} - parent element not found."

        except Exception as e:
            logger.exception(f"Error deleting table in Word document '{file_path}': {e}")
            return f"Error deleting table in Word document '{file_path}': {e}"



if __name__ == '__main__':
    test_file_path = '/Users/xianlindeng/Downloads/乐农〔2025〕69号关于开展2024年度农业"两品"认定奖励申报工作的通知.docx'
    # test_file_path_new = '/Users/xianlindeng/Downloads/test_document_created.docx'
    # test_file_path_modify = '/Users/xianlindeng/Downloads/test_document_modify.docx'
    # test_file_path_table = '/Users/xianlindeng/Downloads/test_document_table.docx'
    # print("--- Testing Create ---")
    # print(CreateWordTool().execute(test_file_path_new))

    # print("\n--- Testing Add Paragraph (Existing File) ---")
    # # Create a file first if it doesn't exist for modification tests
    # if not os.path.exists(test_file_path_modify):
    #     CreateWordTool().execute(test_file_path_modify)
    # print(AddParagraphTool().execute(test_file_path_modify, "This is the first paragraph added.", style="Heading 1"))
    # print(AddParagraphTool().execute(test_file_path_modify, "This is the second paragraph, with default style."))
    # print(AddParagraphTool().execute(test_file_path_modify, "This paragraph tries an invalid style.", style="NonExistentStyle"))

    # print("\n--- Testing Add Table (Existing File) ---")
    # if not os.path.exists(test_file_path_table):
    #      CreateWordTool().execute(test_file_path_table)
    # print(AddTableTool().execute(test_file_path_table, rows=3, cols=4, style="Table Grid", header_row=["ID", "Name", "Value", "Status"]))
    # print(AddTableTool().execute(test_file_path_table, rows=2, cols=2)) # Add another table without style/header

    print("\n--- Testing Read (Original File) ---")
    # Check if the original file exists before trying to open it
    if os.path.exists(test_file_path):
        print(
            ReadWordTool().execute(test_file_path)
        )
    else:
        print(f"Error: Test file not found at '{test_file_path}'")
        print("Please ensure the test file exists or update the path in the script.")

# Update __main__ block for testing (Optional - commented out for safety)
# if __name__ == '__main__':
#     # Example Usage (requires test files in Downloads or adjust paths)
#     base_path = os.path.expanduser("~/Downloads")
#     create_path = os.path.join(base_path, "mrai_test_created.docx")
#     para_path = os.path.join(base_path, "mrai_test_paragraphs.docx")
#     table_path = os.path.join(base_path, "mrai_test_tables.docx")
#     read_path = os.path.join(base_path, "乐农〔2025〕69号关于开展2024年度农业"两品"认定奖励申报工作的通知.docx") # Example existing file

#     # Test Create
#     print("--- Testing Create ---")
#     creator = CreateWordTool()
#     print(creator.execute(create_path))
#     # Clean up created file
#     # if os.path.exists(create_path): os.remove(create_path)

#     # Test Add Paragraph
#     print("\n--- Testing Add Paragraph ---")
#     # Ensure file exists for adding paragraphs
#     if not os.path.exists(para_path):
#         creator.execute(para_path) # Create it first
#     para_adder = AddParagraphTool()
#     print(para_adder.execute(para_path, "First Paragraph - Heading 1", style="Heading 1"))
#     print(para_adder.execute(para_path, "Second Paragraph - Normal Style"))
#     print(para_adder.execute(para_path, "Third Paragraph - Invalid Style", style="FakeStyle123"))

#     # Test Add Table
#     print("\n--- Testing Add Table ---")
#      # Ensure file exists for adding tables
#     if not os.path.exists(table_path):
#         creator.execute(table_path) # Create it first
#     table_adder = AddTableTool()
#     print(table_adder.execute(table_path, rows=4, cols=3, style="Light Shading Accent 1", header_row=["Col A", "Col B", "Col C"]))
#     print(table_adder.execute(table_path, rows=2, cols=2)) # Add a default styled table

#     # Test Read (use an existing file or one just created/modified)
#     print("\n--- Testing Read ---")
#     reader = ReadWordTool()
#     # Try reading the table file we just modified
#     if os.path.exists(table_path):
#         print(f"\nReading: {table_path}")
#         print(reader.execute(table_path))
#     # Try reading the original example file if it exists
#     elif os.path.exists(read_path):
#         print(f"\nReading: {read_path}")
#         print(reader.execute(read_path))
#     else:
#         print(f"\nSkipping Read test - Neither '{table_path}' nor '{read_path}' found.")